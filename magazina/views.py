import base64
import io
from datetime import date, datetime, timedelta
from decimal import Decimal, InvalidOperation
import qrcode
from openpyxl import load_workbook
from django.contrib import messages
from django.contrib.auth.views import LoginView
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.db.models import Count, DecimalField, F, Q, Sum
from django.db.models.functions import Coalesce, ExtractMonth
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from .forms import ExcelImportForm, ProduktiForm, RememberMeAuthenticationForm
from .models import LevizjeStoku, Produkti, has_decimal_part
import pandas as pd
from django.contrib import messages
from .models import Produkti
import json
import openpyxl
from docx import Document
from django.http import HttpResponse
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side



HEADER_ALIASES = {
    'emri': 'emri',
    'produkti': 'emri',
    'produkt': 'emri',
    'pershkrimi': 'pershkrimi',
    'përshkrimi': 'pershkrimi',
    'pershkrim': 'pershkrimi',
    'përshkrim': 'pershkrimi',
    'description': 'pershkrimi',
    'furnitori': 'furnitori',
    'supplier': 'furnitori',
    'sasia': 'sasia',
    'stoku': 'sasia',
    'njesia': 'njesia_matese',
    'njesia_matese': 'njesia_matese',
    'monedha': 'monedha',
    'cmimi_blerjes': 'cmimi_blerjes',
    'çmimi_blerjes': 'cmimi_blerjes',
    'blerja': 'cmimi_blerjes',
    'cmimi_blerje': 'cmimi_blerjes',
    'çmimi_blerje': 'cmimi_blerjes',
    'cmimi_shitjes': 'cmimi_shitjes',
    'çmimi_shitjes': 'cmimi_shitjes',
    'shitja': 'cmimi_shitjes',
    'cmimi_shitje': 'cmimi_shitjes',
    'çmimi_shitje': 'cmimi_shitjes',
    'ne_oferte': 'ne_oferte',
    'oferte': 'ne_oferte',
    'cmimi_ofertes': 'cmimi_ofertes',
    'stoku_minimal': 'stoku_minimal',
    'minimum': 'stoku_minimal',
    'data_skadences': 'data_skadences',
    'skadenca': 'data_skadences',
    'data_hyrjes': 'data_hyrjes',
    'hyrje': 'data_hyrjes',
    'data_daljes': 'data_daljes',
    'dalje': 'data_daljes',
}


REQUIRED_IMPORT_FIELDS = {'emri', 'cmimi_blerjes', 'cmimi_shitjes'}


PRODUCT_SORT_OPTIONS = {
    'emri': ('Emri A-Z', 'emri'),
    '-emri': ('Emri Z-A', '-emri'),
    'sasia': ('Stoku Me I Ulet', 'sasia'),
    '-sasia': ('Stoku Me I Larte', '-sasia'),
    'cmimi_shitjes': ('Cmimi Me I Ulet', 'cmimi_shitjes'),
    '-cmimi_shitjes': ('Cmimi Me I Larte', '-cmimi_shitjes'),
    'data_skadences': ('Skadenca Me E Afert', 'data_skadences'),
    '-krijuar_me': ('Me Te Fundit', '-krijuar_me'),
}


def produktet_e_perdoruesit(user):
    return Produkti.objects.filter(krijuar_nga=user)


def levizjet_e_perdoruesit(user):
    return LevizjeStoku.objects.filter(produkti__krijuar_nga=user)


def parse_date_param(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except ValueError:
        return None


def normalize_header(value):
    return str(value or '').strip().lower().replace(' ', '_').replace('-', '_')


def parse_decimal_value(value, default=None):
    if value in (None, ''):
        return default
    try:
        return Decimal(str(value).replace(',', '.').strip())
    except (InvalidOperation, AttributeError):
        raise ValueError


def parse_bool_value(value):
    if isinstance(value, bool):
        return value
    if value in (None, ''):
        return False
    return str(value).strip().lower() in {'1', 'po', 'yes', 'true', 'y'}


def parse_excel_date(value):
    if value in (None, ''):
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d.%m.%Y'):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except ValueError:
            continue
    raise ValueError


def text_value(value):
    return str(value).strip() if value not in (None, '') else ''


def choice_value(value, allowed_values, default):
    cleaned = text_value(value).upper()
    return cleaned if cleaned in allowed_values else default


def import_products_from_excel(excel_file, user):
    workbook = load_workbook(excel_file, read_only=True, data_only=True)
    sheet = workbook.active
    rows = sheet.iter_rows(values_only=True)
    headers = next(rows, None)

    if not headers:
        return 0, ["File Excel eshte bosh."]

    field_by_column = {}
    for index, header in enumerate(headers):
        field_name = HEADER_ALIASES.get(normalize_header(header))
        if field_name:
            field_by_column[index] = field_name

    missing_fields = REQUIRED_IMPORT_FIELDS - set(field_by_column.values())
    if missing_fields:
        return 0, [f"Mungojne kolonat e detyrueshme: {', '.join(sorted(missing_fields))}."]

    imported = 0
    errors = []

    for row_number, row in enumerate(rows, start=2):
        values = {
            field_name: row[index]
            for index, field_name in field_by_column.items()
            if index < len(row)
        }

        if not any(value not in (None, '') for value in values.values()):
            continue

        try:
            product = Produkti(
                emri=text_value(values.get('emri')),
                pershkrimi=text_value(values.get('pershkrimi')),
                furnitori=text_value(values.get('furnitori')) or None,
                sasia=parse_decimal_value(values.get('sasia'), Decimal('0')),
                njesia_matese=choice_value(values.get('njesia_matese'), {'KG', 'L', 'COPE', 'KUTI'}, 'COPE'),
                monedha=choice_value(values.get('monedha'), {'ALL', 'EUR', 'USD'}, 'ALL'),
                cmimi_blerjes=parse_decimal_value(values.get('cmimi_blerjes')),
                cmimi_shitjes=parse_decimal_value(values.get('cmimi_shitjes')),
                ne_oferte=parse_bool_value(values.get('ne_oferte')),
                cmimi_ofertes=parse_decimal_value(values.get('cmimi_ofertes')),
                stoku_minimal=parse_decimal_value(values.get('stoku_minimal'), Decimal('5')),
                data_skadences=parse_excel_date(values.get('data_skadences')),
                data_hyrjes=parse_excel_date(values.get('data_hyrjes')),
                data_daljes=parse_excel_date(values.get('data_daljes')),
                krijuar_nga=user,
            )
            product.full_clean()
            with transaction.atomic():
                product.save()
                if product.sasia > 0:
                    LevizjeStoku.objects.create(
                        produkti=product,
                        perdoruesi=user,
                        lloji='HYRJE',
                        sasia=product.sasia,
                        vlera=Decimal('0'),
                        stoku_pas=product.sasia,
                    )
            imported += 1
        except Exception as exc:
            errors.append(f"Rreshti {row_number}: {exc}")

    return imported, errors
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except ValueError:
        return None


class RememberMeLoginView(LoginView):
    template_name = 'registration/login.html'
    authentication_form = RememberMeAuthenticationForm

    def form_valid(self, form):
        # 1. Vendosim skadencën e sesionit PËRPARA se Django të bëjë login-in dhe redirect-in
        if form.cleaned_data.get('remember_me'):
            # 60 * 60 * 24 * 30 = 30 ditë (Kodi yt aktual)
            self.request.session.set_expiry(60 * 60 * 24 * 30)
        else:
            # Skadon kur mbyllet browser-i
            self.request.session.set_expiry(0)
            
        # 2. Tani thërrasim super() që të kryejë login-in zyrtar me rregullat e reja të sesionit
        return super().form_valid(form)


@login_required
def lista_produkteve(request):
    query = request.GET.get('q', '').strip()
    njesia = request.GET.get('njesia', '').strip()
    statusi = request.GET.get('statusi', '').strip()
    sort = request.GET.get('sort', 'emri')
    sort_label, sort_field = PRODUCT_SORT_OPTIONS.get(sort, PRODUCT_SORT_OPTIONS['emri'])

    produktet = produktet_e_perdoruesit(request.user)

    if query:
        produktet = produktet.filter(
            Q(emri__icontains=query)
            | Q(pershkrimi__icontains=query)
            | Q(furnitori__icontains=query)
        )

    if njesia in dict(Produkti.NJESIA_CHOICES):
        produktet = produktet.filter(njesia_matese=njesia)

    if statusi == 'oferte':
        produktet = produktet.filter(ne_oferte=True)
    elif statusi == 'stok_ulet':
        produktet = produktet.filter(sasia__lt=F('stoku_minimal'))
    elif statusi == 'skaduar':
        produktet = produktet.filter(data_skadences__lt=timezone.localdate())
    elif statusi == 'skadon_shpejt':
        produktet = produktet.filter(
            data_skadences__gte=timezone.localdate(),
            data_skadences__lte=timezone.localdate() + timedelta(days=7),
        )

    produktet = produktet.order_by(sort_field, 'emri')

    context = {
        'produkte': produktet,
        'query': query,
        'filter_values': {
            'njesia': njesia,
            'statusi': statusi,
            'sort': sort,
        },
        'njesia_choices': Produkti.NJESIA_CHOICES,
        'sort_options': PRODUCT_SORT_OPTIONS,
        'sort_label': sort_label,
        'result_count': produktet.count(),
        'today': timezone.localdate(),
    }
    return render(request, 'magazina/lista.html', context)


@login_required
def skaner(request):
    return render(request, 'magazina/skaner.html')


@login_required
def shto_produkt(request):
    form = ProduktiForm(request.POST or None)

    if request.method == 'POST' and form.is_valid():
        with transaction.atomic():
            produkti = form.save(commit=False)
            produkti.krijuar_nga = request.user
            produkti.save()

            if produkti.sasia > 0:
                LevizjeStoku.objects.create(
                    produkti=produkti,
                    perdoruesi=request.user,
                    lloji='HYRJE',
                    sasia=produkti.sasia,
                    vlera=Decimal('0'),
                    stoku_pas=produkti.sasia,
                )

        messages.success(request, "Produkti u shtua me sukses.")
        return redirect('menaxho_produktin', produkt_id=produkti.id)

    return render(request, 'magazina/shto_produkt.html', {'form': form})


@login_required

def importo_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        file = request.FILES['excel_file']
        
        try:
            # 1. Lexojmë skedarin Excel
            df = pd.read_excel(file)
            
            # 2. Optimizojmë dhe pastrojmë emrat e kolonave
            df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_').str.replace('ç', 'c')
            
            # 3. Fjalori i sinonimeve për kolonat
            perkthimi_kolonave = {
                'emri': ['emri', 'emer', 'produkti', 'emri_i_produktit'],
                'cmimi_blerjes': ['cmimi_blerjes', 'cmimi_i_blerjes', 'blerja', 'cmimi_bleve', 'cmimi_blerje'],
                'cmimi_shitjes': ['cmimi_shitjes', 'cmimi_i_shitjes', 'shitja', 'cmimi_shite', 'cmimi_shitje']
            }
            
            # 4. Ndryshojmë emrat bazuar te sinonimet
            for kolona_django, variacionet in perkthimi_kolonave.items():
                for v in variacionet:
                    if v in df.columns:
                        df.rename(columns={v: kolona_django}, inplace=True)
            
            # 5. Kontrollojmë nëse mungon ndonjë kolonë kritike
            kolonat_e_detyrueshme = ['emri', 'cmimi_blerjes', 'cmimi_shitjes']
            mungojne = [k for k in kolonat_e_detyrueshme if k not in df.columns]
            
            if mungojne:
                messages.error(request, f"Mungojnë kolonat e detyrueshme: {', '.join(mungojne)}")
                return render(request, 'magazina/import.html')
            
            # 6. Fusim të dhënat në Databazë rresht për rresht
            for index, row in df.iterrows():
                Produkti.objects.create(
                    emri=row['emri'],
                    cmimi_blerjes=row['cmimi_blerjes'],
                    cmimi_shitjes=row['cmimi_shitjes'],
                    sasia=row.get('sasia', 0.000),
                    monedha='ALL',               # Vlerë default që të shfaqet në listë
                    njesia_matese='COPE',        # Vlerë default
                    krijuar_nga=request.user     # Lidhja me përdoruesin që bën importin
                )
                
            messages.success(request, "Produktet u importuan me sukses!")
            
        except Exception as e:
            # Ky është blloku 'except' që po kërkonte Python-i
            messages.error(request, f"Gabim gjatë leximit të skedarit: {str(e)}")
            
        return render(request, 'magazina/import.html')
            
    return render(request, 'magazina/import.html')


@login_required
def menaxho_produktin(request, produkt_id):
    produkti = get_object_or_404(produktet_e_perdoruesit(request.user), id=produkt_id)

    if request.method == "POST":
        veprimi = request.POST.get('veprimi')
        shtese_manuale = request.POST.get('shtese_manuale', '').strip()

        try:
            vlera_per_ndryshim = Decimal(shtese_manuale) if shtese_manuale else Decimal('1')
        except InvalidOperation:
            messages.error(request, "Vendos një sasi të vlefshme.")
            return redirect('menaxho_produktin', produkt_id=produkti.id)

        if vlera_per_ndryshim <= 0:
            messages.error(request, "Sasia duhet të jetë më e madhe se zero.")
            return redirect('menaxho_produktin', produkt_id=produkti.id)

        if produkti.njesia_matese != 'KG' and has_decimal_part(vlera_per_ndryshim):
            messages.error(request, "Presja dhjetore lejohet vetëm për produktet me njësi KG.")
            return redirect('menaxho_produktin', produkt_id=produkti.id)

        if veprimi not in {'shto', 'hiq'}:
            messages.error(request, "Veprim i pavlefshëm.")
            return redirect('menaxho_produktin', produkt_id=produkti.id)

        with transaction.atomic():
            produkti = produktet_e_perdoruesit(request.user).select_for_update().get(id=produkt_id)
            if veprimi == "shto":
                produkti.sasia += vlera_per_ndryshim
                lloji_levizjes = 'HYRJE'
            else:
                produkti.sasia = max(produkti.sasia - vlera_per_ndryshim, Decimal('0'))
                lloji_levizjes = 'DALJE'

            produkti.save()
            LevizjeStoku.objects.create(
                produkti=produkti,
                perdoruesi=request.user,
                lloji=lloji_levizjes,
                sasia=vlera_per_ndryshim,
                vlera=vlera_per_ndryshim * produkti.cmimi_shitjes if lloji_levizjes == 'DALJE' else Decimal('0'),
                stoku_pas=produkti.sasia,
            )

        messages.success(request, "Sasia u përditësua me sukses.")
        return redirect('menaxho_produktin', produkt_id=produkti.id)

    stoku_kritik = produkti.sasia < produkti.stoku_minimal
    afer_skadences = False

    if produkti.data_skadences:
        diferenca = produkti.data_skadences - date.today()
        if 0 <= diferenca.days <= 7:
            afer_skadences = True
        elif diferenca.days < 0:
            afer_skadences = "SKADUAR"

    context = {
        'produkti': produkti,
        'stoku_kritik': stoku_kritik,
        'afer_skadences': afer_skadences,
    }

    return render(request, 'magazina/menaxho_produktin.html', context)


@login_required
def dashboard_produkti(request, produkt_id):
    produkti = get_object_or_404(produktet_e_perdoruesit(request.user), id=produkt_id)
    today = timezone.localdate()
    selected_year = request.GET.get('year', str(today.year))

    try:
        selected_year = int(selected_year)
    except ValueError:
        selected_year = today.year

    # Agregimi i lëvizjeve të llojit DALJE nga databaza
    monthly_sales = {
        row['month']: row
        for row in produkti.levizjet.filter(
            lloji='DALJE',
            krijuar_me__year=selected_year,
        ).annotate(
            month=ExtractMonth('krijuar_me')
        ).values('month').annotate(
            total_sasi=Coalesce(Sum('sasia'), Decimal('0'), output_field=DecimalField()),
            total_vlere=Coalesce(Sum('vlera'), Decimal('0'), output_field=DecimalField()),
            nr_levizjesh=Count('id'),
        )
    }

    months = [
        'Janar', 'Shkurt', 'Mars', 'Prill', 'Maj', 'Qershor',
        'Korrik', 'Gusht', 'Shtator', 'Tetor', 'Nëntor', 'Dhjetor',
    ]
    
    # RREDAKTIMI: Krijojmë saktë të gjitha listat që na duhen
    rreshtat_mujore = []
    sasite_list = []
    vlerat_list = []
    total_vjetor_sasi = Decimal('0')
    total_vjetor_vlere = Decimal('0')

    # Ndërtojmë të dhënat muaj pas muaji
    for index, month_name in enumerate(months, start=1):
        row = monthly_sales.get(index, {})
        total_sasi = row.get('total_sasi') or Decimal('0')
        total_vlere = row.get('total_vlere') or Decimal('0')
        
        total_vjetor_sasi += total_sasi
        total_vjetor_vlere += total_vlere
        
        # Ruajmë të dhënat për tabelën tënde ekzistuese në HTML
        rreshtat_mujore.append({
            'muaji': month_name,
            'sasia': total_sasi,
            'vlera': total_vlere,
            'levizje': row.get('nr_levizjesh') or 0,
        })
        
        # Ruajmë të dhënat si numra të thjeshtë (float) për grafikët e Chart.js
        sasite_list.append(float(total_sasi))
        vlerat_list.append(float(total_vlere))

    # Llogaritja e përqindjes për progres-bar-et e tabelës sate
    max_mujor = max((row['sasia'] for row in rreshtat_mujore), default=Decimal('0'))
    for row in rreshtat_mujore:
        row['percent'] = int((row['sasia'] / max_mujor) * 100) if max_mujor else 0

    # Merr 20 lëvizjet e fundit
    levizjet_fundit = produkti.levizjet.select_related('perdoruesi')[:20]

    # Kthejmë çdo gjë të paketuar pastër te template-i HTML
    return render(request, 'magazina/dashboard_produkti.html', {
        'produkti': produkti,
        'selected_year': selected_year,
        'rreshtat_mujore': rreshtat_mujore,
        'total_vjetor_sasi': total_vjetor_sasi,
        'total_vjetor_vlere': total_vjetor_vlere,
        'levizjet_fundit': levizjet_fundit,
        
        # Këto janë listat JSON që kërkojnë grafikët në fund të faqes
        'muajt_json': json.dumps(months),
        'sasite_json': json.dumps(sasite_list),
        'vlerat_json': json.dumps(vlerat_list),
    })
@login_required
def raportet(request):
    today = timezone.localdate()
    period = request.GET.get('period', 'day')
    date_from = parse_date_param(request.GET.get('date_from'))
    date_to = parse_date_param(request.GET.get('date_to'))

    if date_from or date_to:
        start_date = date_from or date_to
        end_date = date_to or date_from
        if start_date > end_date:
            start_date, end_date = end_date, start_date
        active_period = 'custom'
    elif period == 'week':
        start_date = today - timedelta(days=today.weekday())
        end_date = start_date + timedelta(days=6)
        active_period = 'week'
    elif period == 'month':
        start_date = today.replace(day=1)
        next_month = (start_date.replace(day=28) + timedelta(days=4)).replace(day=1)
        end_date = next_month - timedelta(days=1)
        active_period = 'month'
    elif period == 'year':
        start_date = today.replace(month=1, day=1)
        end_date = today.replace(month=12, day=31)
        active_period = 'year'
    else:
        start_date = today
        end_date = today
        active_period = 'day'

    levizjet = levizjet_e_perdoruesit(request.user).select_related('produkti').filter(
        krijuar_me__date__gte=start_date,
        krijuar_me__date__lte=end_date,
    )

    hyrje_rows = levizjet.filter(lloji='HYRJE').values(
        'produkti_id', 'produkti__emri', 'produkti__njesia_matese'
    ).annotate(total=Coalesce(Sum('sasia'), Decimal('0'), output_field=DecimalField()))
    
    # SHTESA KËTU: Marrim edhe 'produkti__cmimi_blerjes' për llogaritjen e kostos
    dalje_rows = levizjet.filter(lloji='DALJE').values(
        'produkti_id', 'produkti__emri', 'produkti__njesia_matese', 'produkti__monedha', 'produkti__cmimi_blerjes'
    ).annotate(
        total=Coalesce(Sum('sasia'), Decimal('0'), output_field=DecimalField()),
        vlere=Coalesce(Sum('vlera'), Decimal('0'), output_field=DecimalField()),
    )

    summary = {}
    for row in hyrje_rows:
        product_id = row['produkti_id']
        summary.setdefault(product_id, {
            'id': product_id,
            'emri': row['produkti__emri'],
            'njesia': row['produkti__njesia_matese'],
            'monedha': '',
            'hyrje': Decimal('0'),
            'dalje': Decimal('0'),
            'vlere_dalje': Decimal('0'),
            'kosto_blerje_dalje': Decimal('0'),  # Vlerë fillestare
            'diferenca_fitimi': Decimal('0'),    # Vlerë fillestare
        })
        summary[product_id]['hyrje'] = row['total']

    for row in dalje_rows:
        product_id = row['produkti_id']
        
        # Llogaritja e kostos së blerjes për sasinë e dalë: Sasia e Shitur * Çmimi i Blerjes
        sasia_shitur = row['total']
        cmimi_blerjes = row['produkti__cmimi_blerjes'] or Decimal('0')
        kosto_blerje = sasia_shitur * cmimi_blerjes
        
        # Diferenca / Fitimi Neto: Vlera e Shitjes - Kosto e Blerjes
        vlera_shitjes = row['vlere']
        diferenca = vlera_shitjes - kosto_blerje

        summary.setdefault(product_id, {
            'id': product_id,
            'emri': row['produkti__emri'],
            'njesia': row['produkti__njesia_matese'],
            'monedha': row['produkti__monedha'],
            'hyrje': Decimal('0'),
            'dalje': Decimal('0'),
            'vlere_dalje': Decimal('0'),
            'kosto_blerje_dalje': Decimal('0'),
            'diferenca_fitimi': Decimal('0'),
        })
        summary[product_id]['dalje'] = sasia_shitur
        summary[product_id]['vlere_dalje'] = vlera_shitjes
        summary[product_id]['monedha'] = row['produkti__monedha']
        summary[product_id]['kosto_blerje_dalje'] = kosto_blerje
        summary[product_id]['diferenca_fitimi'] = diferenca

    # Llogaritja e Totaleve Globale për Kartat e Metrikave
    total_hyrje = sum((row['hyrje'] for row in summary.values()), Decimal('0'))
    total_dalje = sum((row['dalje'] for row in summary.values()), Decimal('0'))
    total_vlere_dalje = sum((row['vlere_dalje'] for row in summary.values()), Decimal('0'))
    
    # SHTESAT E REJA PËR TOTALET FINANCIARE
    total_kosto_blerje = sum((row['kosto_blerje_dalje'] for row in summary.values()), Decimal('0'))
    total_diferenca_fitimi = sum((row['diferenca_fitimi'] for row in summary.values()), Decimal('0'))

    return render(request, 'magazina/raportet.html', {
        'start_date': start_date,
        'end_date': end_date,
        'date_from': request.GET.get('date_from', ''),
        'date_to': request.GET.get('date_to', ''),
        'active_period': active_period,
        'levizjet': levizjet,
        'summary_rows': summary.values(),
        'total_hyrje': total_hyrje,
        'total_dalje': total_dalje,
        'total_vlere_dalje': total_vlere_dalje,
        
        # Kalojmë totalet e reja te template HTML
        'total_kosto_blerje': total_kosto_blerje,
        'total_diferenca_fitimi': total_diferenca_fitimi,
    })

@login_required
def gjenero_qr(request, produkt_id):
    produkti = get_object_or_404(produktet_e_perdoruesit(request.user), id=produkt_id)

    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(str(produkti.id))
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    qr_base64 = base64.b64encode(buffer.getvalue()).decode()

    return render(request, 'magazina/qr_kod.html', {'qr_kod': qr_base64, 'produkti': produkti})



@login_required
def modifiko_produktin(request, produkt_id):
    # 1. Gjejmë produktin specifik nga databaza
    produkt = get_object_or_404(produktet_e_perdoruesit(request.user), id=produkt_id)
    
    if request.method == 'POST':
        # 2. Nëse përdoruesi klikon "Ruaj", përditësojmë të dhënat ekzistuese (instance=produkt)
        forma = ProduktiForm(request.POST, instance=produkt)
        if forma.is_valid():
            forma.save()
            messages.success(request, f"Produkti '{produkt.emri}' u modifikua me sukses!")
            return redirect('home')  # Ose 'lista_produkteve' nëse rruga jote quhet ashtu
    else:
        # 3. Nëse thjesht hapet faqja, mbushim fushat me të dhënat aktuale
        forma = ProduktiForm(instance=produkt)
        
    # Kujdes këtu: Duhet të jetë 'magazina/modifiko_produktin.html' dhe jo importo.html!
    return render(request, 'magazina/modifiko_produktin.html', {'forma': forma, 'produkt': produkt})

@login_required
def eksporto_excel(request):
    # Marrim produktet e përdoruesit loguar
    produktet = Produkti.objects.filter(krijuar_nga=request.user)
    
    # Krijojmë workbook-un e ri të Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Raporti i Inventarit"
    
    # Stilimet bazë për raportin e diplomës
    font_koka = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    fill_koka = PatternFill(start_color="007BFF", end_color="007BFF", fill_type="solid") # Blu zyrtare
    rreshtimi = Alignment(horizontal="center", vertical="center")
    border_holl = Border(
        left=Side(style='thin', color='DDDDDD'), right=Side(style='thin', color='DDDDDD'),
        top=Side(style='thin', color='DDDDDD'), bottom=Side(style='thin', color='DDDDDD')
    )

    # Emrat e kolonave
    headers = ["Produkti", "Furnitori", "Stoku", "Njësia", "Çmimi Blerjes", "Çmimi Shitjes", "Monedha", "Vlera Totale"]
    ws.append(headers)
    
    # Apliko stilin te koka e tabelës
    for cell in ws[1]:
        cell.font = font_koka
        cell.fill = fill_koka
        cell.alignment = rreshtimi

    # Plotësimi i të dhënave nga databaza
    for p in produktet:
        vlera_totale = float(p.sasia) * float(p.cmimi_shitjes)
        row = [
            p.emri,
            p.furnitori or "-",
            float(p.sasia),
            p.njesia_matese,
            float(p.cmimi_blerjes),
            float(p.cmimi_shitjes),
            p.monedha,
            vlera_totale
        ]
        ws.append(row)
    
    # Formatimi i numrave me dy shifra pas presjes (.00) fiks siç e kërkove
    for row in ws.iter_rows(min_row=2, max_row=len(produktet)+1):
        for cell in row:
            cell.border = border_holl
            # Nëse jemi te kolona e sasisë, çmimeve ose vlerës totale, vendosim formatin dhjetor
            if cell.column in [3, 5, 6, 8]:
                cell.number_format = '#,##0.00'

    # Rregullimi automatik i gjerësisë së kolonave
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

    # Përgatitja e përgjigjes për shkarkim
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = f'attachment; filename=Raporti_Inventarit_{timezone.localdate()}.xlsx'
    wb.save(response)
    return response


# ==========================================
# 2. EKSPORTI NË WORD (.docx)
# ==========================================
@login_required
def eksporto_word(request):
    produktet = Produkti.objects.filter(krijuar_nga=request.user)
    
    doc = Document()
    doc.add_heading('RAPORTI ZYRTAR I INVENTARIT', level=1)
    doc.add_paragraph(f"Gjeneruar më: {timezone.now().strftime('%d/%m/%Y %H:%M')}")
    
    # Krijojmë tabelën në Word
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1' # Stili i gatshëm profesional i Word
    
    # Koka e tabelës
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Produkti'
    hdr_cells[1].text = 'Furnitori'
    hdr_cells[2].text = 'Stoku'
    hdr_cells[3].text = 'Çmimi Blerjes'
    hdr_cells[4].text = 'Çmimi Shitjes'
    hdr_cells[5].text = 'Vlera'
    
    # Hedhja e të dhënave
    for p in produktet:
        row_cells = table.add_row().cells
        vlera = float(p.sasia) * float(p.cmimi_shitjes)
        
        row_cells[0].text = str(p.emri)
        row_cells[1].text = str(p.furnitori or "-")
        row_cells[2].text = f"{float(p.sasia):.2f} {p.njesia_matese}"
        row_cells[3].text = f"{float(p.cmimi_blerjes):.2f} {p.monedha}"
        row_cells[4].text = f"{float(p.cmimi_shitjes):.2f} {p.monedha}"
        row_cells[5].text = f"{vlera:.2f} {p.monedha}"
        
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Raporti_Inventarit_{timezone.localdate()}.docx'
    doc.save(response)
    return response


from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from .models import Produkti

def fshi_produktin(request, produkt_id):
    produkti = get_object_or_404(Produkti, id=produkt_id)
    if produkti.qr_code:
        produkti.qr_code.delete(save=False)
    emri = produkti.emri
    produkti.delete()
    messages.success(request, f"Produkti '{emri}' u fshi me sukses!")
    return redirect('lista_produkteve')  # Ose rruga ku kthehet lista e produkteve